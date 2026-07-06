import sys
import tempfile
from pathlib import Path
from uuid import uuid4

import httpx
from docx import Document


BASE_URL = "http://127.0.0.1:8000"


def check(response: httpx.Response, expected: int = 200):
    if response.status_code != expected:
        raise RuntimeError(f"{response.request.method} {response.request.url}: {response.status_code} {response.text}")
    return response.json()


def headers(token: str) -> dict:
    return {"Authorization": f"Bearer {token}"}


def main():
    suffix = uuid4().hex[:8]
    client = httpx.Client(base_url=BASE_URL, timeout=60)

    admin_login = check(client.post("/auth/login", json={"email": "admin@smartedu.com", "password": "admin123"}))
    admin_token = admin_login["access_token"]
    assert admin_login["user"]["role"] == "admin"
    assert check(client.get("/auth/me", headers=headers(admin_token)))["email"] == "admin@smartedu.com"

    faculty_email = f"faculty-{suffix}@example.com"
    faculty = check(client.post("/auth/create_faculty", headers=headers(admin_token), json={
        "name": "Smoke Faculty", "email": faculty_email, "password": "faculty123"
    }))
    faculty_login = check(client.post("/auth/login", json={"email": faculty_email, "password": "faculty123"}))
    faculty_token = faculty_login["access_token"]

    student_email = f"student-{suffix}@example.com"
    student_register = check(client.post("/auth/register", json={
        "name": "Smoke Student", "email": student_email, "password": "student123"
    }))
    student_token = student_register["access_token"]
    forbidden = client.post("/subjects", headers=headers(student_token), json={"name": "Blocked", "code": f"B{suffix}", "description": ""})
    assert forbidden.status_code == 403

    subject = check(client.post("/subjects", headers=headers(faculty_token), json={
        "name": f"Computer Networks {suffix}",
        "code": f"CN{suffix}",
        "description": "Networking fundamentals",
    }))
    subject_id = subject["id"]

    with tempfile.TemporaryDirectory() as directory:
        path = Path(directory) / "network-notes.docx"
        doc = Document()
        doc.add_heading("Computer Networks", 0)
        doc.add_paragraph(
            "The Transmission Control Protocol, known as TCP, provides reliable, ordered, "
            "and error-checked delivery of data between applications. TCP uses acknowledgements "
            "and retransmission to recover from packet loss."
        )
        doc.add_paragraph(
            "The Internet Protocol routes packets across networks. Unlike TCP, IP itself does "
            "not guarantee delivery, ordering, or duplicate protection."
        )
        doc.save(path)
        with path.open("rb") as stream:
            upload = check(client.post(
                f"/materials/upload?subject_id={subject_id}",
                headers=headers(faculty_token),
                files={"file": (path.name, stream, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            ))
    assert upload["chunks"] >= 1
    assert check(client.post(f"/subjects/{subject_id}/rebuild", headers=headers(faculty_token)))["chunks"] >= 1

    materials = check(client.get(f"/materials?subject_id={subject_id}", headers=headers(student_token)))
    assert materials[0]["filename"] == "network-notes.docx"

    answer = check(client.post("/rag/ask", headers=headers(student_token), json={
        "subject_id": subject_id,
        "question": "How does TCP recover from packet loss?",
    }))
    assert answer["document_name"] == "network-notes.docx"
    assert "retransmission" in answer["answer"].lower()
    assert answer["page_number"] == 1

    quiz = check(client.post("/quiz/generate", headers=headers(student_token), json={
        "subject_id": subject_id, "difficulty": "medium", "question_count": 5,
    }))
    assert len(quiz["questions"]) == 5
    summary = check(client.post("/summary/generate", headers=headers(student_token), json={"subject_id": subject_id}))
    assert summary["summary"]
    assert check(client.get("/rag/history", headers=headers(student_token)))[0]["document_name"] == "network-notes.docx"

    analytics = check(client.get("/admin/analytics", headers=headers(admin_token)))
    assert analytics["total_documents"] >= 1 and analytics["total_questions_asked"] >= 1

    check(client.delete(f"/subjects/{subject_id}", headers=headers(admin_token)))
    check(client.delete(f"/admin/users/{faculty['id']}", headers=headers(admin_token)))
    check(client.delete(f"/admin/users/{student_register['user']['id']}", headers=headers(admin_token)))
    check(client.post("/auth/logout", headers=headers(admin_token)))
    print("PASS: auth, roles, admin bootstrap, subjects, DOCX upload, chunks, FAISS, RAG citation, quiz, summary, history, analytics")


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(f"FAIL: {exc}", file=sys.stderr)
        raise

